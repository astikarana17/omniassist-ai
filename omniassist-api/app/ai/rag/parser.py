"""Document parsing: PDF, DOCX, TXT/MD, and HTML (for crawled pages)."""
from __future__ import annotations

import io

from app.core.exceptions import ValidationError
from app.core.logging import get_logger

logger = get_logger("rag.parser")

SUPPORTED = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "txt",
    "text/html": "html",
}


def parse(content: bytes, content_type: str, filename: str = "") -> str:
    """Extract plain text from a document. Raises ValidationError for unsupported types."""
    kind = SUPPORTED.get(content_type)
    if kind is None:
        # Fall back on extension when content-type is generic.
        lower = filename.lower()
        if lower.endswith(".pdf"):
            kind = "pdf"
        elif lower.endswith(".docx"):
            kind = "docx"
        elif lower.endswith((".txt", ".md")):
            kind = "txt"
        elif lower.endswith((".html", ".htm")):
            kind = "html"
        else:
            raise ValidationError(
                f"Unsupported file type: {content_type or filename}", code="UNSUPPORTED_FILE"
            )

    if kind == "pdf":
        return _parse_pdf(content)
    if kind == "docx":
        return _parse_docx(content)
    if kind == "html":
        return _parse_html(content)
    return content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return _normalize("\n\n".join(pages))


def _parse_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return _normalize("\n".join(parts))


def _parse_html(content: bytes) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()
    return _normalize(soup.get_text(separator="\n"))


def _normalize(text: str) -> str:
    lines = [ln.strip() for ln in text.splitlines()]
    cleaned = [ln for ln in lines if ln]
    return "\n".join(cleaned).strip()
